"""Document parsing service for PDF and DOCX resumes."""
import io
from typing import Optional, Dict, Any, Tuple
from pathlib import Path

import pdfplumber
from docx import Document


class DocumentParsingError(Exception):
    """Exception raised when document parsing fails."""
    pass


class DocumentParser:
    """Service for extracting raw text and layout metadata from PDF/DOCX files.
    
    Note: This service only handles document extraction. Structural parsing
    (experiences, education, bullets) should be done by a separate Parser Agent.
    """
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc'}
    
    def __init__(self):
        """Initialize the document parser."""
        pass
    
    def extract_raw_text(
        self, 
        file_path: Optional[str] = None,
        file_bytes: Optional[bytes] = None,
        file_extension: Optional[str] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract raw text and layout metadata from a resume file.
        
        Args:
            file_path: Path to the resume file
            file_bytes: Raw bytes of the file
            file_extension: File extension (required if using file_bytes)
            
        Returns:
            Tuple of (raw_text, layout_metadata)
            
        Raises:
            DocumentParsingError: If extraction fails
        """
        if file_path:
            path = Path(file_path)
            extension = path.suffix.lower()
            
            if extension not in self.SUPPORTED_FORMATS:
                raise DocumentParsingError(
                    f"Unsupported file format: {extension}. "
                    f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
                )
            
            with open(path, 'rb') as f:
                file_bytes = f.read()
        
        elif file_bytes and file_extension:
            extension = file_extension.lower()
            if not extension.startswith('.'):
                extension = f'.{extension}'
        else:
            raise DocumentParsingError("Must provide either file_path or (file_bytes + file_extension)")
        
        # Extract text based on format
        if extension == '.pdf':
            text = self._extract_pdf_text(file_bytes)
            layout_metadata = self._extract_pdf_layout(file_bytes)
        elif extension in {'.docx', '.doc'}:
            text = self._extract_docx_text(file_bytes)
            layout_metadata = self._extract_docx_layout(file_bytes)
        else:
            raise DocumentParsingError(f"Unsupported extension: {extension}")
        
        return text, layout_metadata
    
    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        """Extract text from PDF using pdfplumber."""
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return '\n\n'.join(text_parts)
        except Exception as e:
            raise DocumentParsingError(f"Failed to extract PDF text: {str(e)}")
    
    def _extract_pdf_layout(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract layout metadata from PDF."""
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                layout = {
                    'page_count': len(pdf.pages),
                    'page_dimensions': [],
                    'fonts': set(),
                }
                
                for page in pdf.pages:
                    layout['page_dimensions'].append({
                        'width': page.width,
                        'height': page.height
                    })
                    
                    # Extract font info if available
                    if hasattr(page, 'chars'):
                        for char in page.chars:
                            if 'fontname' in char:
                                layout['fonts'].add(char['fontname'])
                
                layout['fonts'] = list(layout['fonts'])
                return layout
        except Exception as e:
            return {'error': str(e)}
    
    def _extract_docx_text(self, file_bytes: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise DocumentParsingError(f"Failed to extract DOCX text: {str(e)}")
    
    def _extract_docx_layout(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract layout metadata from DOCX."""
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            layout = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'sections': len(doc.sections),
                'styles': set(),
            }
            
            # Extract style information
            for paragraph in doc.paragraphs:
                if paragraph.style:
                    layout['styles'].add(paragraph.style.name)
            
            layout['styles'] = list(layout['styles'])
            return layout
        except Exception as e:
            return {'error': str(e)}
